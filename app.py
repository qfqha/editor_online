import os
import json
import uuid
import pandas as pd
import docx
from io import BytesIO
from flask import Flask, render_template, request, redirect, url_for, session, jsonify, send_file
from flask_socketio import SocketIO, emit, join_room, leave_room

# 初始化Flask应用
app = Flask(__name__)
app.config['SECRET_KEY'] = 'collaborative-editor-secret-key'  # 生产环境需更换
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['ALLOWED_EXTENSIONS'] = {'docx', 'xlsx', 'xls'}

# 确保上传目录存在
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# 初始化SocketIO
socketio = SocketIO(app, cors_allowed_origins="*")

# 存储用户数据（测试用）
users = {
    "admin": "admin123",
    "user1": "password1",
    "user2": "password2"
}

# 存储文件数据和编辑状态
file_data = {}  # file_id -> {id, name, type, content, editors, path}
online_users = set()

# 辅助函数：检查文件扩展名
def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

# 路由：登录页面
@app.route('/', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        
        if username in users and users[username] == password:
            session['username'] = username
            online_users.add(username)
            # 通知所有用户在线状态更新
            socketio.emit('online_users_updated', list(online_users))
            return redirect(url_for('file_list'))
        else:
            return render_template('login.html', error="用户名或密码错误")
    
    if 'username' in session:
        return redirect(url_for('file_list'))
        
    return render_template('login.html')

# 路由：登出
@app.route('/logout')
def logout():
    if 'username' in session:
        username = session['username']
        if username in online_users:
            online_users.remove(username)
        session.pop('username', None)
        # 通知所有用户在线状态更新
        socketio.emit('online_users_updated', list(online_users))
    return redirect(url_for('login'))

# 路由：文件列表
@app.route('/files')
def file_list():
    if 'username' not in session:
        return redirect(url_for('login'))
        
    return render_template('index.html', 
                         files=list(file_data.values()),
                         online_users=online_users)

# 路由：上传文件
@app.route('/upload', methods=['POST'])
def upload_file():
    if 'username' not in session:
        return jsonify({"status": "error", "message": "请先登录"})
    
    if 'file' not in request.files:
        return jsonify({"status": "error", "message": "没有文件部分"})
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({"status": "error", "message": "未选择文件"})
    
    if file and allowed_file(file.filename):
        file_ext = file.filename.rsplit('.', 1)[1].lower()
        file_id = str(uuid.uuid4())
        filename = f"{file_id}.{file_ext}"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        content = ""
        file_type = "word" if file_ext == "docx" else "excel"
        
        try:
            if file_type == "word":
                doc = docx.Document(filepath)
                paragraphs = [{'text': p.text, 'style': p.style.name} for p in doc.paragraphs]
                content = json.dumps(paragraphs)
            else:  # excel
                df = pd.read_excel(filepath)
                content = df.to_json(orient="split")
                
            file_data[file_id] = {
                "id": file_id,
                "name": file.filename,
                "type": file_type,
                "content": content,
                "editors": {},  # 确保是字典类型
                "path": filepath
            }
            
            return jsonify({
                "status": "success", 
                "message": "文件上传成功",
                "file_id": file_id
            })
        except Exception as e:
            return jsonify({"status": "error", "message": f"文件处理错误: {str(e)}"})
    
    return jsonify({"status": "error", "message": "不支持的文件格式"})

# 路由：编辑文件
@app.route('/edit/<file_id>')
def edit_file(file_id):
    if 'username' not in session:
        return redirect(url_for('login'))
    
    if file_id not in file_data:
        return "文件不存在", 404
    
    # 确保editors是字典
    if not isinstance(file_data[file_id].get('editors', {}), dict):
        file_data[file_id]['editors'] = {}
        
    return render_template('editor.html', 
                         file=file_data[file_id],
                         username=session['username'])

# 路由：下载文件
@app.route('/download/<file_id>')
def download_file(file_id):
    if 'username' not in session:
        return redirect(url_for('login'))
    
    if file_id not in file_data:
        return "文件不存在", 404
    
    file_info = file_data[file_id]
    try:
        if file_info['type'] == 'word':
            # 从JSON内容创建Word文档
            doc = docx.Document()
            paragraphs = json.loads(file_info['content'])
            for para in paragraphs:
                p = doc.add_paragraph(para.get('text', ''))
                # 尝试应用样式
                if 'style' in para:
                    try:
                        p.style = para['style']
                    except:
                        pass  # 忽略无效样式
            
            # 保存到内存
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            return send_file(
                buffer,
                as_attachment=True,
                download_name=file_info['name'],
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            
        elif file_info['type'] == 'excel':
            # 从JSON内容创建Excel文件
            data = json.loads(file_info['content'])
            df = pd.DataFrame(
                data['data'], 
                columns=data['columns'],
                index=data.get('index')
            )
            
            # 保存到内存
            buffer = BytesIO()
            with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
                df.to_excel(writer, index=False)
            
            buffer.seek(0)
            
            return send_file(
                buffer,
                as_attachment=True,
                download_name=file_info['name'],
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            
    except Exception as e:
        return f"文件生成错误: {str(e)}", 500

# SocketIO事件：加入编辑
@socketio.on('join_editor')
def handle_join_editor(file_id):
    if 'username' not in session:
        return
    
    username = session['username']
    join_room(file_id)
    
    if file_id in file_data:
        file_data[file_id]['editors'][username] = "文档开始处"
        emit('user_joined_editor', {
            'username': username,
            'editors': file_data[file_id]['editors']
        }, room=file_id, include_self=False)
    
    if file_id in file_data:
        emit('current_content', {
            'content': file_data[file_id]['content'],
            'editors': file_data[file_id]['editors']
        })

# SocketIO事件：处理编辑动作
@socketio.on('editor_action')
def handle_editor_action(data):
    if 'username' not in session:
        return
    
    file_id = data['file_id']
    if file_id not in file_data:
        return
    
    username = session['username']
    # 更新文件内容和编辑位置
    file_data[file_id]['content'] = data['content']
    file_data[file_id]['editors'][username] = data.get('position', '未指定位置')
    
    # 广播给房间内其他用户
    emit('content_updated', {
        'content': data['content'],
        'editors': file_data[file_id]['editors'],
        'username': username
    }, room=file_id, include_self=False)
    
    # 向发送者确认同步完成
    emit('sync_complete')

# SocketIO事件：处理Excel结构变更（增删行列）
@socketio.on('excel_structure_change')
def handle_excel_structure_change(data):
    if 'username' not in session:
        return
    
    file_id = data['file_id']
    if file_id not in file_data:
        return
    
    username = session['username']
    # 更新文件内容
    file_data[file_id]['content'] = data['content']
    file_data[file_id]['editors'][username] = data.get('position', '修改了表格结构')
    
    # 广播结构变更给房间内其他用户
    emit('excel_structure_updated', {
        'content': data['content'],
        'editors': file_data[file_id]['editors'],
        'action': data['action'],
        'username': username
    }, room=file_id, include_self=False)
    
    # 向发送者确认同步完成
    emit('sync_complete')

# SocketIO事件：离开编辑
@socketio.on('leave_editor')
def handle_leave_editor(file_id):
    if 'username' not in session:
        return
    
    username = session['username']
    leave_room(file_id)
    
    if file_id in file_data and username in file_data[file_id]['editors']:
        del file_data[file_id]['editors'][username]
        emit('user_left_editor', {
            'username': username,
            'editors': file_data[file_id]['editors']
        }, room=file_id)

# 路由：获取在线用户
@app.route('/online-users')
def get_online_users():
    return jsonify(list(online_users))

if __name__ == '__main__':
    socketio.run(app, host='0.0.0.0', port=5000, debug=True, use_reloader=False)
    

